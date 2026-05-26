from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "PokerVision_Project_Milestone.docx"
FIGURE = ROOT / "outputs" / "demo" / "annotated" / "frame_0004.png"


def bold_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    return paragraph


def labeled_paragraph(doc: Document, label: str, value: str):
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value)
    return paragraph


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_grid_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
    return table


def main():
    doc = Document()

    bold_paragraph(doc, "PokerVision Project Milestone")
    doc.add_paragraph("Automated hand-history reconstruction from broadcast-style poker frames")
    labeled_paragraph(doc, "Student", "Ameya Jadhav")
    labeled_paragraph(doc, "Project", "PokerVision: Automated Hand History Reconstruction from Poker Videos")
    labeled_paragraph(doc, "Milestone", "Component 2: Project Milestone")
    labeled_paragraph(doc, "Date", "May 26, 2026")

    bold_paragraph(doc, "Technical Progress")
    doc.add_paragraph(
        "The implementation now covers the main pipeline promised in the proposal: frame ingestion, fixed-layout region extraction, card recognition, overlay text recognition, temporal tracking, and structured hand-history output. The current version is packaged as a command-line tool with a reproducible synthetic broadcast demo and JSON outputs for both frame-level detections and reconstructed hand events."
    )
    doc.add_paragraph(
        "ROI processing is config-driven for community cards, player cards, pot text, and action text. Card recognition uses template matching over a 52-card deck and returns both the card label and confidence score for each visible slot. The tracker converts repeated visual detections into a timeline of streets, pot changes, revealed hole cards, and player actions."
    )

    bold_paragraph(doc, "Visualization and Intermediate Results")
    doc.add_paragraph(
        "Figure 1 shows the annotated river frame from the current demo run. Yellow boxes mark community-card ROIs, green boxes mark player-card ROIs, and cyan boxes mark text overlays. The final reconstructed state is board 7H 8D 2C QS AD, Alice KD QD, Bob AH TS, final pot 390."
    )
    if FIGURE.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = paragraph.add_run().add_picture(str(FIGURE), width=Inches(5.5))
        picture._inline.docPr.set(
            "descr",
            "Annotated PokerVision demo frame with ROI boxes around community cards, player cards, pot text, and action text.",
        )
    doc.add_paragraph("Figure 1. Annotated ROI detections on a synthetic broadcast river frame.")

    doc.add_page_break()

    bold_paragraph(doc, "Results Snapshot")
    add_grid_table(
        doc,
        ["Component", "Milestone result", "Current status"],
        [
            ("Card recognition", "Template matcher over a 52-card deck; ROI crops identify visible board and player cards.", "Demo: all visible cards recovered"),
            ("Overlay OCR", "Glyph OCR plus candidate matching for fixed pot/action overlays.", "Demo: 5 pot updates and 5 actions parsed"),
            ("State tracking", "Finite-state timeline records hole cards, streets, pot deltas, and actions.", "Demo: 15 structured events"),
        ],
    )
    doc.add_paragraph(
        "The demo produces 15 timeline events: two hole-card reveals, five pot updates, five parsed actions, and flop/turn/river transitions. This confirms that the system is functioning as an end-to-end prototype rather than only a collection of isolated detectors."
    )

    bold_paragraph(doc, "Deviations from Original Proposal")
    doc.add_paragraph(
        "The main pivot is that the current milestone uses a controlled synthetic broadcast layout before moving to YouTube footage. This keeps the work aligned with the original proposal's fixed-layout assumption while removing dataset and broadcast-style variability during pipeline construction. The core technical pieces remain the same: ROI extraction, template matching, OCR, tracking, and finite-state reconstruction. Real-video ingestion is supported as an optional OpenCV path, but the next step is to tune templates and ROIs on sampled frames from one actual broadcast."
    )
    bold_paragraph(doc, "Current limitations")
    doc.add_paragraph(
        "The prototype still assumes stable camera graphics, known text locations, and a small action vocabulary. Real broadcasts will require cropped card templates from the target overlay, validation labels, and smoothing for noisy OCR or partially occluded cards."
    )

    bold_paragraph(doc, "Updated Timeline and Objectives")
    add_grid_table(
        doc,
        ["Remaining phase", "Objectives", "Deliverable"],
        [
            ("Week 3", "Calibrate ROIs on one real broadcast layout; collect sampled frames; add cropped broadcast card templates; label a small validation set.", "Real-frame run"),
            ("Week 3-4", "Evaluate card accuracy, overlay OCR accuracy, and event-level correctness; add smoothing for repeated/ambiguous detections.", "Metrics table"),
            ("Week 4", "Generate annotated demo clip, finalize hand-history JSON examples, and write the final report with failure analysis.", "Final submission"),
        ],
    )
    doc.add_paragraph(
        "Final objective: demonstrate PokerVision on a short clip from one consistent poker broadcast format, report detector-level accuracy and event-level reconstruction accuracy, and include an annotated video plus structured hand-history output."
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
